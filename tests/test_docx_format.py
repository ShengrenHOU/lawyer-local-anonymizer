import base64
import zipfile

from docx import Document

from legal_anonymizer.document_io import read_text_document, scan_docx_unsupported_parts
from legal_anonymizer.pipeline import anonymize_file, restore_file_auto
from legal_anonymizer.workspace import create_workspace


def test_docx_anonymize_and_restore_preserves_word_container_format(tmp_path):
    workspace = create_workspace(tmp_path)
    source = workspace.pending / "demo.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "联系电话：13800000000"
    paragraph = document.add_paragraph()
    run = paragraph.add_run("乙方：张三")
    run.bold = True
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "上海测试科技有限公司"
    document.save(source)

    anonymized = anonymize_file(source, workspace)

    assert anonymized.output_path.suffix == ".docx"
    anonymized_doc = Document(anonymized.output_path)
    assert len(anonymized_doc.tables) == 1
    assert anonymized_doc.paragraphs[0].runs[0].bold is True
    anonymized_text = read_text_document(anonymized.output_path)
    assert "张三" not in anonymized_text
    assert "13800000000" not in anonymized_text
    assert "[[PERSON_001]]" in anonymized_text
    assert "[[PHONE_001]]" in anonymized_text

    restored = restore_file_auto(anonymized.output_path, workspace)

    assert restored.output_path.suffix == ".docx"
    restored_doc = Document(restored.output_path)
    assert len(restored_doc.tables) == 1
    assert restored_doc.paragraphs[0].runs[0].bold is True
    restored_text = read_text_document(restored.output_path)
    assert "张三" in restored_text
    assert "13800000000" in restored_text


def test_docx_anonymize_collapses_company_suffix_split_across_runs(tmp_path):
    workspace = create_workspace(tmp_path)
    source = workspace.pending / "split-company.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Goodwish")
    paragraph.add_run(" International Limited")
    document.save(source)

    anonymized = anonymize_file(source, workspace)

    anonymized_text = read_text_document(anonymized.output_path)
    assert "International Limited" not in anonymized_text
    assert "[[COMPANY_" in anonymized_text


def test_docx_replacement_preserves_unrelated_run_styles(tmp_path):
    workspace = create_workspace(tmp_path)
    source = workspace.pending / "mixed-style.docx"
    document = Document()
    paragraph = document.add_paragraph()
    label = paragraph.add_run("FROM: ")
    label.bold = True
    name = paragraph.add_run("Alice Chen")
    name.italic = True
    tail = paragraph.add_run(" signs here.")
    tail.underline = True
    document.save(source)

    anonymized = anonymize_file(source, workspace)

    anonymized_doc = Document(anonymized.output_path)
    runs = anonymized_doc.paragraphs[0].runs
    assert runs[0].text == "FROM: "
    assert runs[0].bold is True
    assert "[[PERSON_" in runs[1].text
    assert runs[1].italic is True
    assert runs[2].text == " signs here."
    assert runs[2].underline is True


def test_docx_reader_includes_additional_word_xml_parts(tmp_path):
    source = tmp_path / "comments.docx"
    document = Document()
    document.add_paragraph("Body text")
    document.save(source)
    with zipfile.ZipFile(source, "a", zipfile.ZIP_DEFLATED) as package:
        package.writestr(
            "word/comments.xml",
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="0"><w:p><w:r><w:t>Secret Comment Limited</w:t></w:r></w:p></w:comment>
</w:comments>""",
        )

    text = read_text_document(source)

    assert "Secret Comment Limited" in text


def test_docx_with_image_is_forced_to_review(tmp_path):
    workspace = create_workspace(tmp_path)
    source = workspace.pending / "with-image.docx"
    image_path = tmp_path / "tiny.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
    )
    document = Document()
    document.add_paragraph("Rockit Trading (Shanghai) Co., Ltd.")
    document.add_picture(str(image_path))
    document.save(source)

    result = anonymize_file(source, workspace)
    structural_findings = scan_docx_unsupported_parts(source)

    assert structural_findings
    assert result.output_path.parent == workspace.review_required
    assert not result.upload_allowed
    assert any(finding.category == "UNSUPPORTED_IMAGE" for finding in result.risk_findings)
