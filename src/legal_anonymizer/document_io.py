from __future__ import annotations

import tempfile
import zipfile
import re
from html import escape
from pathlib import Path
from typing import Iterable

import fitz
from docx import Document
from docx.document import Document as DocxDocument
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from legal_anonymizer.models import MappingTable


def read_text_document(path: Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8")
    if suffix == ".docx":
        document = Document(str(path))
        return "\n".join(paragraph.text for paragraph in _iter_docx_paragraphs(document))
    if suffix == ".doc":
        with tempfile.TemporaryDirectory(prefix="legal-anonymizer-doc-") as temp_dir:
            docx_path = Path(temp_dir) / f"{path.stem}.docx"
            _convert_doc_to_docx(path, docx_path)
            document = Document(str(docx_path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
    if suffix == ".pdf":
        with fitz.open(str(path)) as pdf:
            return "\n".join(page.get_text("text") for page in pdf)
    raise ValueError(f"Unsupported document type: {suffix}")


def write_text_document(path: Path, text: str) -> Path:
    path = Path(path)
    suffix = path.suffix.lower()
    path.parent.mkdir(parents=True, exist_ok=True)
    if suffix in {".txt", ".md"}:
        path.write_text(text, encoding="utf-8")
        return path
    if suffix == ".docx":
        document = Document()
        for line in text.splitlines() or [""]:
            document.add_paragraph(line)
        document.save(str(path))
        return path
    raise ValueError(f"Unsupported output type: {suffix}")


def anonymize_docx_document(source: Path, target: Path, table: MappingTable) -> Path:
    replacements = {mapping.value: mapping.placeholder for mapping in table.mappings}
    return _replace_docx_document(source, target, replacements)


def restore_docx_document(source: Path, target: Path, table: MappingTable) -> Path:
    replacements = {mapping.placeholder: mapping.value for mapping in table.mappings}
    return _replace_docx_document(source, target, replacements)


def _replace_docx_document(source: Path, target: Path, replacements: dict[str, str]) -> Path:
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document(str(source))
    ordered = dict(sorted(replacements.items(), key=lambda item: -len(item[0])))
    for paragraph in _iter_docx_paragraphs(document):
        _replace_paragraph_text(paragraph, ordered)
    document.save(str(target))
    _replace_docx_xml_text(target, ordered)
    return target


def _iter_docx_paragraphs(document: DocxDocument) -> Iterable[Paragraph]:
    yield from document.paragraphs
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        for container in (section.header, section.footer):
            yield from container.paragraphs
            for table in container.tables:
                yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterable[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            yield from _iter_cell_paragraphs(cell)


def _iter_cell_paragraphs(cell: _Cell) -> Iterable[Paragraph]:
    yield from cell.paragraphs
    for table in cell.tables:
        yield from _iter_table_paragraphs(table)


def _replace_paragraph_text(paragraph: Paragraph, replacements: dict[str, str]) -> None:
    if not paragraph.text:
        return
    for run in paragraph.runs:
        run.text = _replace_text(run.text, replacements)
    collapsed = _collapse_placeholder_company_suffix(paragraph.text)
    if any(source in paragraph.text for source in replacements) or collapsed != paragraph.text:
        replaced = _collapse_placeholder_company_suffix(_replace_text(paragraph.text, replacements))
        if paragraph.runs:
            paragraph.runs[0].text = replaced
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(replaced)


def _replace_text(text: str, replacements: dict[str, str]) -> str:
    result = text
    for source, target in replacements.items():
        if source:
            result = result.replace(source, target)
    return result


def _collapse_placeholder_company_suffix(text: str) -> str:
    return re.sub(
        r"(\[\[COMPANY_\d{3}\]\])\s+"
        r"(?:International\s+)?(?:Limited|Ltd\.?|LLC|LLP|Inc\.?|Corporation|Company Limited)\b",
        r"\1",
        text,
        flags=re.IGNORECASE,
    )


def _replace_docx_xml_text(path: Path, replacements: dict[str, str]) -> None:
    temp_path = path.with_suffix(f"{path.suffix}.tmp")
    with zipfile.ZipFile(path, "r") as source_zip, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for item in source_zip.infolist():
            data = source_zip.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                text = data.decode("utf-8")
                for source, target in replacements.items():
                    text = text.replace(escape(source, quote=False), escape(target, quote=False))
                data = text.encode("utf-8")
            target_zip.writestr(item, data)
    path.write_bytes(temp_path.read_bytes())
    temp_path.unlink(missing_ok=True)


def _convert_doc_to_docx(source: Path, target: Path) -> None:
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise RuntimeError(
            "当前文件是旧版 .doc 格式。请安装 Microsoft Word 后重试，"
            "或先用 Word/WPS 另存为 .docx 再放入待匿名化文件夹。"
        ) from exc

    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(source.resolve()), ReadOnly=True)
        document.SaveAs2(str(target.resolve()), FileFormat=16)
    except Exception as exc:
        raise RuntimeError(
            "无法自动读取旧版 .doc 文件。请确认本机已安装 Microsoft Word，"
            "或先将文件另存为 .docx 后再处理。"
        ) from exc
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()
